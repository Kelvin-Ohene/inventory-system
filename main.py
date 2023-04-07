from flask import Flask, render_template, request, redirect,url_for, session
import pymysql
import mysql.connector
import re
from tkinter import *
import tkinter as tk
from PIL import ImageTk, Image  
import re
from tkinter import messagebox
from tkinter import ttk
import mysql.connector
import re
import pymysql, os
import credentials as cr
import socket
from datetime import date
from mysql.connector import connect, Error
import pdfkit
import pandas as pd
import csv
import docx

app = Flask(__name__)
app.secret_key = 'secretkey'


# Establish a connection to the MySQL server
cnx = mysql.connector.connect(
  host="#",
  database="#",
  user="#",
  password="#"
)
cur = cnx.cursor(buffered=True)

regex = '^[a-z0-9]+[\._]?[a-z0-9]+[@]\w+[.]\w{2,3}$'

def validate(u_input):
    if (re.search(regex, u_input) and u_input.isalpha):
        #print("Correct")
        return "Correct"
    else:
        #print("Wrong")
        return "Wrong"

def only_numbers(char):
    return char.isdigit()


@app.route('/')
def inventory_form():
  url = "http://192.168.10.18:5052/dashboard"
  table = pd.read_html(url)[0]
  table.to_excel("dashboard.xlsx") 
  return render_template('login2.html')

@app.route('/login', methods=['POST'])
def login():
  username = request.form.get('username')
  password = request.form.get('password')
  
  try:
      conn = mysql.connector.connect(host="#", user="#", password="#")
      cur = conn.cursor(buffered=True)
      cur.execute("use inventory")
      cur.execute("select * from login_info where name=%s and password=%s",(username,password))
      row=cur.fetchone()
      if row == None:
        print("Error!","Invalid USERNAME & PASSWORD")
        error_code = "Error!","Invalid USERNAME & PASSWORD"
        return render_template('login.html', error_code = error_code)
      else:
        print("here")
        session['username'] = username
        return render_template('index.html')
        
          
                    
                

  except Exception as e:
          messagebox.showerror("Error!",f"Error due to {str(e)}")
  #win = Toplevel()
  #infopage1.info(win)
  
  """if authenticated:
    # Login the user and redirect to the dashboard
    login_user()
    return redirect(url_for('dashboard'))
  else:
    # Show an error message
    return render_template('login.html', error='Invalid username or password')"""


@app.route('/login-to-registration')
def login_to_registration():
  return render_template("registration.html")


@app.route('/registration', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        # Get form data
        username = request.form['username']
        email = request.form['email']
        password = request.form['password']
        confirm_password = request.form['confirm_password']

        # Validate form data
        if password != confirm_password:
            return render_template('registration.html', error='Passwords do not match')
        else:
          if validate(email) == "Wrong":
            #print("Error", "Invalid E-mail.")
            return render_template('registration.html', error='Wrong E-mail format')
          elif email == "" or password == "" or confirm_password == "" or username == "":
              #print("Error", "All fields must be filled")
              return render_template('register.html', error='All fields empty')
          else:
              cur.execute(f"insert into login_info(name,email,password, confirmPassword) values ('{username}','{email}','{password}','{confirm_password}')")
              cnx.commit()
              
              return render_template('inventory.html', success='Registration successful')
          return render_template('login2.html')


@app.route('/submit_item', methods=['POST'])
def submit_item():
  # Extract data from the request payload
  item_received_from = request.form.get('item_received_from')
  item_contact_person = request.form.get('item_contact_person')
  item_contact = request.form.get('item_contact')
  item_email = request.form.get('item_email')
  item_company = request.form.get('item_company')
  item_manufacturer = request.form.get('item_manufacturer')
  item_equipment = request.form.get('item_equipment')
  item_model = request.form.get('item_model')
  item_serial_number = request.form.get('item_serial_number')
  item_warranty_status = request.form.get('item_warranty_status')
  item_warranty_info = request.form.get('item_warranty_info')
  item_complaint = request.form.get('item_complaint')
  item_date = request.form.get('item_date')
  item_place_of_service = request.form.get('radio')
  item_current_location = request.form.get('item_current_location')
  item_technician = request.form.get('item_technician')
  item_update = request.form.get('item_update')
  print(item_technician)
  
  if not item_received_from or not item_technician:
        # Return a JavaScript alert to the user
        return '<script>alert("Field Empty");</script>'
        
  else:
        cursor = cnx.cursor()
        # Insert data into the database
        cursor.execute("use inventory")
        cursor.execute(f"insert into inventory_items(name, contact_person, contact, email, company, manufacturer, product, model, serial_no, warranty, warranty_info , complaint, datee, place_of_service, current_location, technician , progress) values ('{item_received_from}','{item_contact_person}','{item_contact}','{item_email}','{item_company}','{item_manufacturer}','{item_equipment}','{item_model}','{item_serial_number}','{item_warranty_status}','{item_warranty_info}','{item_complaint}','{item_date}','{item_place_of_service}','{item_current_location}','{item_technician}','{item_update}')")
        cnx.commit()
        success = "Data entered successfully"
        return render_template('inventory.html', success = success)
  
  # Validate the data
  #if not item_received_from or not item_technician:
  #  return "Item name and quantity are required!", 400
  # Save the data to the database or write it to a file
  # Return a response to the client
  #return "Item added to inventory successfully!"


@app.route('/dashboard',methods=['POST','GET'])
def show_dashboard():

    query = 'SELECT * FROM inventory_items'
    cur.execute(query)
    data = cur.fetchall()
    #print(data)
    return render_template('dashboard.html', data=data)
  
def connect_db():
    connection = pymysql.connect(host='#',
                             user='#',
                             password='#',
                             db='#',
                             charset='utf8mb4',
                             cursorclass=pymysql.cursors.DictCursor)
    return connection  
  
@app.route('/search', methods=['POST','GET'])
def serach():
  return render_template('search.html')


@app.route('/search_results', methods=['POST'])
def search_results():
    search_term = request.form['search_term']
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            # Search for item
            sql = "SELECT * FROM inventory_items WHERE name LIKE %s OR id LIKE %s"
            cursor.execute(f"SELECT * FROM inventory_items WHERE company='{search_term}';")
            results = cursor.fetchall()
            print(results)
            print("Done")

        # close the connection
        connection.close()
        return render_template('search_results.html', results=results)

    except Exception as e:
        return(str(e))

@app.route('/update', methods=['POST'])
def update():
    item_id = request.form['item_id']
    new_quantity = request.form['new_quantity']

    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            # update the item's quantity
            sql = "UPDATE inventory_items SET progress = concat(progress, %s '.','\n') WHERE id = %s"
            cursor.execute(sql, (new_quantity, item_id))
            connection.commit()
            cnx.commit()
            #cnx.close()
            #CustomerName=concat(CustomerName,",Brown")

        # close the connection
        connection.close()
        return render_template('inventory.html',inserted="Quantity updated successfully!")

    except Exception as e:
        return(str(e))

@app.route('/logout')
def logout():
    session.pop('username', None) # remove the username from the session
    return redirect('/')
  
@app.route('/to-inventory')
def toInventory():
    return render_template("inventory.html")
  
@app.route('/index')
def index():
  connection = connect_db()
  try:
      with connection.cursor() as cursor:
          # Search for item
          today = date.today()
          sql = "SELECT * FROM inventory_items WHERE name LIKE %s OR id LIKE %s"
          #cursor.execute(f"SELECT * FROM inventory_items WHERE company='{search_term}';")
          cursor.execute(f"SELECT * FROM inventory_items WHERE DATE(datee) = '{today}';")
          results = cursor.fetchall()
          print(results)
          print("Done Clicked")
          print(today)
          

      # close the connection
      connection.close()
      return render_template('index.html', results=results)

  except Exception as e:
      return(str(e))
    
@app.route('/to_submit_data_testing')
def to_submit_data_testing(): 
  return render_template("testing.html")
    
@app.route('/submit_data_testing', methods=['POST'])
def submit_data_testing():
    try:
        connection = connect(
            host='#',
            user='#',
            password='#',
            database='#'
        )
        
        cursor = connection.cursor()
        address = request.form['address']
        name = request.form['name']
        telephone_number = request.form['telephone_number']
        contact = request.form['contact']
        email = request.form['email']
        complaint = request.form['complaint']
        work_supervisor = request.form['work_supervisor']
        date_started = request.form['date_started']
        date_completed = request.form['date_completed']
        equipment = request.form['equipment']
        serial_number = request.form['serial_number']
        model = request.form['model']
        technician = request.form['technician']
        technicians_report = request.form['technicians_report']
        recommendation = request.form['recommendation']
        
        cursor.execute(f"INSERT INTO testing (address, name, telephone_number, contact, email, complaint, work_supervisor, date_started, date_completed, equipment, serial_number, model, technician, technicians_report, recommendation) VALUES ('{address}', '{name}', '{telephone_number}', '{contact}', '{email}', '{complaint}', '{work_supervisor}', '{date_started}', '{date_completed}', '{equipment}', '{serial_number}', '{model}', '{technician}', '{technicians_report}', '{recommendation}');")
        connection.commit()
        cursor.close()
        connection.close()
        return "Data submitted successfully."
    except Error as e:
        return str(e)
  
@app.route('/create_table', methods=['GET'])
def create_table():
    try:
        connection = connect(
            host='#',
            user='#',
            password='#',
            database='#'
        )
        
        cursor = connection.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS testing (
                address VARCHAR(255),
                name VARCHAR(255),
                telephone_number VARCHAR(255),
                contact VARCHAR(255),
                email VARCHAR(255),
                complaint VARCHAR(255),
                work_supervisor VARCHAR(255),
                date_started VARCHAR(255),
                date_completed VARCHAR(255),
                equipment VARCHAR(255),
                serial_number VARCHAR(255),
                model VARCHAR(255),
                technician VARCHAR(255),
                technicians_report VARCHAR(255),
                recommendation VARCHAR(255)
            );
        """)
        connection.commit()
        cursor.close()
        connection.close()
        return "Table 'testing' created successfully."
    except Error as e:
        return str(e)

@app.route('/testing_dashboard',methods=['POST','GET'])
def show_testing_dashboard():

    query = 'SELECT * FROM testing'
    cur.execute(query)
    data = cur.fetchall()
    #print(data)
    return render_template('testing_dashboard.html', data=data)

@app.route('/hmtl-to-pdf')
def html_to_pdf():
  pdfkit.from_url('https://www.google.com/','sample.pdf') 
  
@app.route('/to-update')
def to_update():
    return render_template("update_form.html")

@app.route('/to-table',methods=['POST','GET'])
def to_table():
    if request.method == "POST":
        # Create the Word document
        doc = docx.Document()

        # Save the document
        doc.save("Report.docx")
        address = request.form.get("address")
        equipment = request.form.get("equipment")
        serial_number = request.form.get("serial_number")
        fault = request.form.get("fault")
        current_location = request.form.get("current_location")
        status = request.form.get("status")
        update = request.form.get("update")

        doc = docx.Document("Report.docx")
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = "Address"
        header_cells[1].text = "Equipment"
        header_cells[2].text = "Serial Number"
        header_cells[3].text = "Fault"
        header_cells[4].text = "Current Location"
        header_cells[5].text = "Status"
        header_cells[6].text = "Update"

        new_row = table.add_row().cells
        new_row[0].text = address
        new_row[1].text = equipment
        new_row[2].text = serial_number
        new_row[3].text = fault
        new_row[4].text = current_location
        new_row[5].text = status
        new_row[6].text = update

        doc.save("Report.docx")

        return "Data submitted successfully!"


if __name__ == '__main__':
  #inventory = {}
  app.run(host="192.168.10.18",port=5052)

